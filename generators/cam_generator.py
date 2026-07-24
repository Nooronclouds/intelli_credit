import json
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paths ────────────────────────────────────────────────────────────
FINANCIAL_DATA_PATH = "outputs/financial_data.json"
RESEARCH_DATA_PATH  = "outputs/research_data.json"
DECISION_PATH       = "outputs/decision.json"
CAM_OUTPUT          = "outputs/cam.docx"


# ── Styling helpers ──────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Set background color of a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1, color="1F3864"):
    """Add a styled heading paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(13) if level == 1 else Pt(11)
    run.font.color.rgb = RGBColor(
        int(color[0:2], 16),
        int(color[2:4], 16),
        int(color[4:6], 16)
    )
    # Bottom border on heading
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_kv_table(doc, rows, col_widths=(2.5, 4.0)):
    """Add a simple two-column key-value table."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = 'Table Grid'
    for i, (key, val) in enumerate(rows):
        cells = table.rows[i].cells
        cells[0].width = Inches(col_widths[0])
        cells[1].width = Inches(col_widths[1])

        k_run = cells[0].paragraphs[0].add_run(str(key))
        k_run.bold      = True
        k_run.font.size = Pt(9)
        set_cell_bg(cells[0], "EBF3FB")

        v_run = cells[1].paragraphs[0].add_run(str(val) if val is not None else "N/A")
        v_run.font.size = Pt(9)
    doc.add_paragraph()
    return table


def add_financial_table(doc, headers, data_rows, header_color="1F3864"):
    """Add a multi-column financial table with styled header."""
    table = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr_cells[i], header_color)
        p   = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold           = True
        run.font.size      = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        p.alignment        = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for ri, row in enumerate(data_rows):
        cells = table.rows[ri + 1].cells
        bg    = "F2F9FF" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            set_cell_bg(cells[ci], bg)
            p   = cells[ci].paragraphs[0]
            run = p.add_run(str(val) if val is not None else "N/A")
            run.font.size = Pt(9)
            if ci > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()
    return table


def format_cr(val):
    """Format a crore value cleanly."""
    if val is None:
        return "N/A"
    return f"Rs. {val:,.2f} Cr"


def format_ratio(val, suffix="x"):
    """Format a ratio value cleanly."""
    if val is None:
        return "N/A"
    return f"{val:.2f}{suffix}"


def decision_color(decision):
    """Return hex color for decision badge."""
    return {
        "APPROVE": "1E8449",
        "REVIEW":  "B7950B",
        "REJECT":  "922B21",
    }.get(decision, "555555")


# ── CAM Generator ────────────────────────────────────────────────────

def generate_cam(financial_data=None, research_data=None, decision_data=None):
    """
    MAIN FUNCTION — generates a Credit Approval Memo Word document.
    Reads from outputs/ if data not passed directly.
    """
    print(f"\n{'='*50}")
    print("CAM GENERATOR — Credit Approval Memo")
    print(f"{'='*50}")

    # Load data
    if financial_data is None and os.path.exists(FINANCIAL_DATA_PATH):
        with open(FINANCIAL_DATA_PATH) as f:
            financial_data = json.load(f)

    if research_data is None and os.path.exists(RESEARCH_DATA_PATH):
        with open(RESEARCH_DATA_PATH) as f:
            research_data = json.load(f)

    if decision_data is None and os.path.exists(DECISION_PATH):
        with open(DECISION_PATH) as f:
            decision_data = json.load(f)

    if not financial_data or not decision_data:
        print("  ERROR: Missing required data files. Run ingestor and recommender first.")
        return None

    company    = financial_data.get("company", {})
    financials = financial_data.get("financials", {})
    gst        = financial_data.get("gst_analysis", {})
    ct         = financial_data.get("circular_trading", {})
    bank       = financial_data.get("bank_analysis", {})
    promoter   = research_data.get("promoter_risk", {}) if research_data else {}
    litigation = research_data.get("litigation_risk", {}) if research_data else {}
    sector     = research_data.get("sector_risk", {}) if research_data else {}
    red_flags  = decision_data.get("red_flags", [])
    score      = decision_data.get("credit_score", 0)
    decision   = decision_data.get("decision", "REVIEW")
    rationale  = decision_data.get("decision_rationale", "")
    breakdown  = decision_data.get("score_breakdown", {})

    doc = Document()

    # ── Page margins ─────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Title block ──────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("CREDIT APPROVAL MEMORANDUM (CAM)")
    title_run.bold           = True
    title_run.font.size      = Pt(16)
    title_run.font.color.rgb = RGBColor(31, 56, 100)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(
        f"Prepared by Intelli-Credit AI System  |  {datetime.now().strftime('%d %B %Y')}"
    )
    sub_run.font.size      = Pt(9)
    sub_run.font.color.rgb = RGBColor(100, 100, 100)
    sub_run.italic         = True

    doc.add_paragraph()

    # ── Decision banner ──────────────────────────────────────────────
    banner_table = doc.add_table(rows=1, cols=3)
    banner_table.style = 'Table Grid'
    dcol = decision_color(decision)

    cells = banner_table.rows[0].cells
    # Company name
    set_cell_bg(cells[0], "EBF3FB")
    p = cells[0].paragraphs[0]
    r = p.add_run(company.get("name", "Unknown Company"))
    r.bold = True; r.font.size = Pt(12)

    # Score
    set_cell_bg(cells[1], "F8F9FA")
    p = cells[1].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Credit Score\n{score}/100")
    r.bold = True; r.font.size = Pt(12)

    # Decision
    set_cell_bg(cells[2], dcol)
    p = cells[2].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(decision)
    r.bold = True; r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph()

    # ── 1. Company Profile ───────────────────────────────────────────
    add_heading(doc, "1. COMPANY PROFILE")
    add_kv_table(doc, [
        ("Company Name",    company.get("name")),
        ("CIN",             company.get("cin")),
        ("PAN",             company.get("pan")),
        ("Industry",        company.get("industry")),
        ("Loan Requested",  format_cr(company.get("loan_requested_cr"))),
        ("Directors",       ", ".join(company.get("directors", [])) or "N/A"),
        ("Financial Year",  str(financials.get("year", "N/A"))),
    ])

    # ── 2. Financial Summary ─────────────────────────────────────────
    add_heading(doc, "2. FINANCIAL SUMMARY")
    doc.add_paragraph().add_run(
        "All figures in Indian Rupees Crores (Rs. Cr) unless stated otherwise."
    ).font.size = Pt(8)

    add_financial_table(doc,
        headers=["Metric", "Value", "Assessment"],
        data_rows=[
            ["Revenue from Operations",    format_cr(financials.get("revenue_cr")),
             "Primary income stream"],
            ["Net Profit (PAT)",           format_cr(financials.get("net_profit_cr")),
             "Positive" if (financials.get("net_profit_cr") or 0) > 0 else "Loss-making"],
            ["Total Assets",               format_cr(financials.get("total_assets_cr")),
             "Asset base"],
            ["Total Debt / Borrowings",    format_cr(financials.get("total_debt_cr")),
             "Debt obligations"],
            ["Equity / Net Worth",         format_cr(financials.get("equity_cr")),
             "Shareholder value"],
            ["EBITDA",                     format_cr(financials.get("ebitda_cr")),
             "Operating profitability"],
            ["Cash & Equivalents",         format_cr(financials.get("cash_cr")),
             "Liquidity buffer"],
        ]
    )

    # ── 3. Key Financial Ratios ──────────────────────────────────────
    add_heading(doc, "3. KEY FINANCIAL RATIOS")

    dscr  = financials.get("dscr")
    dte   = financials.get("debt_to_equity")
    cr    = financials.get("current_ratio")
    ccr   = financials.get("collateral_coverage_ratio")

    def ratio_status(val, good_threshold, bad_threshold, higher_is_better=True):
        if val is None:
            return "N/A"
        if higher_is_better:
            return "STRONG" if val >= good_threshold else ("ADEQUATE" if val >= bad_threshold else "WEAK")
        else:
            return "STRONG" if val <= good_threshold else ("ADEQUATE" if val <= bad_threshold else "WEAK")

    add_financial_table(doc,
        headers=["Ratio", "Value", "Benchmark", "Status"],
        data_rows=[
            ["DSCR (Debt Service Coverage)",
             format_ratio(dscr), ">= 1.5x",
             ratio_status(dscr, 2.0, 1.5)],
            ["Debt to Equity",
             format_ratio(dte), "<= 2.0x",
             ratio_status(dte, 1.0, 2.0, higher_is_better=False)],
            ["Current Ratio",
             format_ratio(cr), ">= 1.2x",
             ratio_status(cr, 1.5, 1.2)],
            ["Collateral Coverage",
             format_ratio(ccr), ">= 1.5x",
             ratio_status(ccr, 2.0, 1.5)],
        ]
    )

    # ── 4. GST & Fraud Analysis ──────────────────────────────────────
    add_heading(doc, "4. GST RECONCILIATION & FRAUD SIGNALS")

    gst_flag = gst.get("flag", "UNKNOWN")
    gst_color = {"CLEAN": "1E8449", "MODERATE_MISMATCH": "B7950B",
                 "REVENUE_INFLATION": "922B21"}.get(gst_flag, "555555")

    add_financial_table(doc,
        headers=["Check", "Finding", "Risk Level", "Score Impact"],
        data_rows=[
            ["GSTR-3B Revenue",
             format_cr(gst.get("gstr3b_revenue_cr")),
             "Self-declared", ""],
            ["GSTR-2A Revenue",
             format_cr(gst.get("gstr2a_revenue_cr")),
             "Supplier-declared", ""],
            ["GST Mismatch",
             f"{gst.get('mismatch_percentage') or 0:.1f}%",
             gst_flag,
             f"{gst.get('score_impact', 0):+d} pts"],
            ["Circular Trading",
             f"{ct.get('cycles_found', 0)} cycle(s) detected",
             ct.get("risk_level", "LOW"),
             f"{ct.get('score_impact', 0):+d} pts"],
        ]
    )

    if gst.get("note"):
        p = doc.add_paragraph()
        r = p.add_run(f"Note: {gst.get('note')}")
        r.italic = True; r.font.size = Pt(8)
    doc.add_paragraph()

    # ── 5. Bank Statement Analysis ───────────────────────────────────
    add_heading(doc, "5. BANK STATEMENT ANALYSIS")
    add_kv_table(doc, [
        ("Avg Monthly Credits",  format_cr(bank.get("avg_monthly_credit_cr"))),
        ("Avg Monthly Debits",   format_cr(bank.get("avg_monthly_debit_cr"))),
        ("Minimum Balance",      format_cr(bank.get("min_balance_cr"))),
        ("Bounce Count (6 mths)", str(bank.get("bounce_count_6months", 0))),
        ("Irregular Patterns",   str(len(bank.get("irregular_patterns", []))) + " detected"),
    ])

    # ── 6. Research & Risk Intelligence ─────────────────────────────
    add_heading(doc, "6. RESEARCH & RISK INTELLIGENCE")

    add_financial_table(doc,
        headers=["Risk Category", "Finding", "Score Impact"],
        data_rows=[
            ["Promoter / Adverse News",
             "Adverse news detected" if promoter.get("adverse_news_found") else "No adverse news",
             f"{promoter.get('score_impact', 0):+d} pts"],
            ["Litigation",
             f"{litigation.get('cases_found', 0)} pending case(s)",
             f"{litigation.get('score_impact', 0):+d} pts"],
            [f"Sector Risk ({sector.get('sector', 'N/A')})",
             sector.get("sector_sentiment", "N/A"),
             f"{sector.get('score_impact', 0):+d} pts"],
        ]
    )

    # Sector headwinds/tailwinds
    if sector.get("headwinds") or sector.get("tailwinds"):
        p = doc.add_paragraph()
        p.add_run("Sector Headwinds: ").bold = True
        p.add_run(", ".join(sector.get("headwinds", ["None"])))
        p.runs[-1].font.size = Pt(9)

        p2 = doc.add_paragraph()
        p2.add_run("Sector Tailwinds: ").bold = True
        p2.add_run(", ".join(sector.get("tailwinds", ["None"])))
        p2.runs[-1].font.size = Pt(9)

    doc.add_paragraph()

    # ── 7. Red Flags ─────────────────────────────────────────────────
    add_heading(doc, "7. RED FLAGS SUMMARY", color="922B21")

    if red_flags:
        add_financial_table(doc,
            headers=["Category", "Description", "Severity", "Score Impact"],
            data_rows=[
                [f["category"], f["description"], f["severity"],
                 f"{f.get('score_impact', 0):+d} pts"]
                for f in red_flags
            ],
            header_color="922B21"
        )
    else:
        p = doc.add_paragraph()
        p.add_run("No critical red flags identified.").font.size = Pt(9)
        doc.add_paragraph()

    # ── 8. Score Breakdown ───────────────────────────────────────────
    add_heading(doc, "8. CREDIT SCORE BREAKDOWN")
    add_financial_table(doc,
        headers=["Component", "Points"],
        data_rows=[
            ["Base Score",                       f"{breakdown.get('base_score', 100):+d}"],
            ["Financial Health Score",            f"{breakdown.get('financial_score', 0):+d}"],
            ["Fraud Signal Penalties",            f"{breakdown.get('fraud_signal_score', 0):+d}"],
            ["Research Risk Penalties",           f"{breakdown.get('research_score', 0):+d}"],
            ["FINAL CREDIT SCORE",                f"{breakdown.get('final_score', 0)} / 100"],
        ]
    )

    # ── 9. Credit Decision ───────────────────────────────────────────
    add_heading(doc, "9. CREDIT DECISION & RECOMMENDATION")

    decision_para = doc.add_paragraph()
    decision_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = decision_para.add_run(f"DECISION: {decision}")
    dr.bold = True; dr.font.size = Pt(14)
    dr.font.color.rgb = RGBColor(
        int(decision_color(decision)[0:2], 16),
        int(decision_color(decision)[2:4], 16),
        int(decision_color(decision)[4:6], 16)
    )

    doc.add_paragraph()
    rationale_para = doc.add_paragraph()
    rationale_para.add_run("Rationale: ").bold = True
    rationale_para.add_run(rationale)

    doc.add_paragraph()

    if decision == "APPROVE":
        rec_text = (
            "The applicant demonstrates adequate financial health with acceptable risk levels. "
            "Credit may be extended subject to standard documentation and compliance checks. "
            "Regular monitoring of GST filings and annual financial statements is recommended."
        )
    elif decision == "REVIEW":
        rec_text = (
            "The application falls within the review band and requires assessment by a senior "
            "credit officer. Key concerns should be addressed through additional documentation, "
            "collateral verification, and management discussion before a final decision is made."
        )
    else:
        rec_text = (
            "The application does not meet minimum credit standards based on current financial "
            "data and risk assessment. Credit should not be extended at this time. The applicant "
            "may reapply after addressing the identified risk factors."
        )

    rec_para = doc.add_paragraph()
    rec_para.add_run("Recommendation: ").bold = True
    rec_para.add_run(rec_text).font.size = Pt(9)

    # ── 10. Disclaimer ───────────────────────────────────────────────
    doc.add_paragraph()
    add_heading(doc, "DISCLAIMER", color="888888")
    disclaimer = doc.add_paragraph()
    dr = disclaimer.add_run(
        "This Credit Approval Memorandum has been generated by the Intelli-Credit AI system "
        "using automated analysis of submitted financial documents, GST data, bank statements, "
        "and publicly available information. This document is intended to assist credit officers "
        "and does not constitute a final credit decision. All credit decisions must be reviewed "
        "and approved by authorized personnel in accordance with the institution's credit policy."
    )
    dr.font.size = Pt(8)
    dr.italic    = True
    dr.font.color.rgb = RGBColor(120, 120, 120)

    # Save
    os.makedirs("outputs", exist_ok=True)
    # Close any existing file before overwriting
    if os.path.exists(CAM_OUTPUT):
        try:
            os.remove(CAM_OUTPUT)
        except PermissionError:
            print("  ERROR: cam.docx is open in Word — please close it and run again.")
            return None
    doc.save(CAM_OUTPUT)

    print(f"\n✓ cam.docx saved to {CAM_OUTPUT}")
    print(f"{'='*50}\n")
    return CAM_OUTPUT


if __name__ == "__main__":
    generate_cam()